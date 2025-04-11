from typing import List
from llama_index.core import Document
from llama_index.core.readers.base import BaseReader
import os

class ExcelReader(BaseReader):
    """Process Excel files into LlamaIndex documents."""
    def load_data(self, file_path: str, extra_info=None) -> List[Document]:
        try:
            import pandas as pd
            str_path = str(file_path)  # Convert PosixPath to string
            df = pd.read_excel(str_path)
            text = df.to_string()
            return [Document(
                text=text,
                metadata={
                    "file_name": os.path.basename(str_path),
                    "file_path": str_path,
                    "file_type": "excel"
                }
            )]
        except ImportError:
            raise ValueError("Excel support requires pandas. Install with: pip install pandas")

class DocxReader(BaseReader):
    """Process Word documents into LlamaIndex documents."""
    def load_data(self, file_path: str, extra_info=None) -> List[Document]:
        try:
            from docx import Document as DocxDocument
            str_path = str(file_path)  # Convert PosixPath to string
            doc = DocxDocument(str_path)
            # Extract text from paragraphs
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
            return [Document(
                text=text,
                metadata={
                    "file_name": os.path.basename(str_path),
                    "file_path": str_path,
                    "file_type": "docx"
                }
            )]
        except ImportError:
            raise ValueError("Word document support requires python-docx. Install with: pip install python-docx")
